import subprocess
from rest_framework import viewsets
from .models import Trip, DailyLog
from .serializers import TripSerializer, DailyLogSerializer
from django.http import JsonResponse

import matplotlib
matplotlib.use('Agg')

import matplotlib.pyplot as plt
import numpy as np
import io
import base64
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import tempfile
import os
import zipfile
import sys



class TripViewSet(viewsets.ModelViewSet):
    queryset = Trip.objects.all()
    serializer_class = TripSerializer

class DailyLogViewSet(viewsets.ModelViewSet):
    queryset = DailyLog.objects.all()
    serializer_class = DailyLogSerializer

def get_libreoffice_executable():
    if sys.platform.startswith('linux'):
        # Check common Linux paths
        paths = [
            '/usr/bin/libreoffice',
            '/usr/local/bin/libreoffice',
        ]
        for path in paths:
            if os.path.exists(path):
                return path
    elif sys.platform == 'darwin':  # macOS
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'

    return None  # LibreOffice not found


def generate_trip_pdfs(request, trip_id):
    # Fetch all DailyLogs associated with the given trip_id
    daily_logs = DailyLog.objects.filter(tripId=trip_id)  #  Fixed Foreign Key reference

    if not daily_logs.exists():
        return HttpResponse("No daily logs found for the given trip ID", status=404)

    temp_dir = tempfile.mkdtemp()  # Create a temporary directory for files
    pdf_paths = []  # Store paths of generated PDFs

    for daily_log in daily_logs:
        dailylogId = daily_log.dailylogId
        trip = daily_log.tripId  #  Fixed Foreign Key reference

        driver_name = trip.driver_name
        date = daily_log.date.strftime("%Y-%m-%d")
        total_miles_driven = daily_log.total_miles_driven
        total_mileage = daily_log.total_mileage
        carrier_name = daily_log.carrier_name
        vehicle_details = daily_log.vehicle_details
        duty_status_data = daily_log.duty_status  # Extract stored JSON field
        from_address = daily_log.from_address
        to_address = daily_log.to_address
        home_terminal_address = daily_log.home_terminal_address
        main_office_address = daily_log.main_office_address
        remarks = daily_log.remarks

        # **Generate Chart**
        status_mapping = {"onDuty": -0.5, "driving": 0.5,  "sleeperBerth": 1.5,  "offDuty": 2.5  }
        activity_labels = ["On Duty", "Driving", "Sleeper Berth", "Off Duty"]
        colors = ["#F9DCC4", "#F4A261", "#264653", "#2A9D8F"]

        time = []
        activity = []
        for entry in duty_status_data:
            time.append(entry["start_time"])
            activity.append(status_mapping[entry["status"]])
            time.append(entry["end_time"])
            activity.append(status_mapping[entry["status"]])

        fig, ax = plt.subplots(figsize=(12, 4))
        for i in range(4):  
            ax.fill_between([0, 24], i-1, i, color=colors[i], alpha=0.5)

        ax.step(time, activity, where='post', color='black', linewidth=2)
        ax.set_yticks([0, 1, 2, 3])
        ax.set_yticklabels(activity_labels, fontsize=12, weight="bold")
        ax.set_xticks(np.arange(0, 25, 1))
        ax.xaxis.set_label_position("top")
        ax.xaxis.tick_top()
        ax.set_xlabel("Time (Hours)", fontsize=12, weight="bold")
        ax.set_ylabel("Activity", fontsize=12, weight="bold")
        ax.set_title(f"Driver's Daily Log (ID: {dailylogId})")
        ax.grid(True, linestyle="--", alpha=0.6)

        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        plt.close(fig)
        buffer.seek(0)

        # **Load Existing .docx Template (Use Django Settings)**
        from django.conf import settings
        doc_template_path = os.path.join(settings.BASE_DIR, "trips/tripLogGenerater.docx")
        doc = Document(doc_template_path)  #  Fixed Hardcoded Path

        replacements = {
            "DRIVER_NAME": str(driver_name) if driver_name else "",
            "DATE": date,
            "TOTAL_MILES_DRIVEN": str(total_miles_driven) if total_miles_driven else "0",
            "TOTAL_MILEAGE": str(total_mileage) if total_mileage else "0",
            "CARRIER_NAME": str(carrier_name) if carrier_name else "",
            "VEHICLE_DETAILS": str(vehicle_details) if vehicle_details else "",
            "FROM_ADDRESS": str(from_address) if from_address else "",
            "TO_ADDRESS": str(to_address) if to_address else "",
            "MAIN_OFFICE_ADDRESS": str(main_office_address) if main_office_address else "",
            "HOME_TERMINAL_ADDRESS": str(home_terminal_address) if home_terminal_address else "",
            "ANY_REMARKS": str(remarks) if remarks else ""
        }

        # **Replace Text Placeholders**
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # **Replace Table Cells**
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # **Replace GENERATED_IMAGE Placeholder with Image**
        for paragraph in doc.paragraphs:
            if "GENERATED_IMAGE" in paragraph.text:
                paragraph.text = paragraph.text.replace("GENERATED_IMAGE", "")
                run = paragraph.add_run()
                image_stream = io.BytesIO(buffer.getvalue())
                run.add_picture(image_stream, width=Inches(6))

        # **Save .docx Temporarily**
        temp_docx_path = os.path.join(temp_dir, f"daily_log_{dailylogId}.docx")
        doc.save(temp_docx_path)

        # **Convert .docx to PDF**
        output_pdf_path = temp_docx_path.replace(".docx", ".pdf")
   
        libreoffice_exec = get_libreoffice_executable()
        try:
            subprocess.run([
                libreoffice_exec,
                "--headless", "--convert-to", "pdf", temp_docx_path, "--outdir", temp_dir
            ], check=True)
            pdf_paths.append(output_pdf_path)
        except subprocess.CalledProcessError as e:
            return HttpResponse(f" PDF Conversion Failed: {e}", status=500)

    # **Create a ZIP file containing all PDFs**
    zip_path = os.path.join(temp_dir, f"trip_{trip_id}_dailylogs.zip")
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for pdf_path in pdf_paths:
            zipf.write(pdf_path, os.path.basename(pdf_path))

    # **Read ZIP and Return as Response**
    with open(zip_path, "rb") as zip_file:
        response = HttpResponse(zip_file.read(), content_type="application/zip")
        response["Content-Disposition"] = f'attachment; filename="trip_{trip_id}_dailylogs.zip"'

    # Cleanup temporary files
    for file_path in pdf_paths:
        os.remove(file_path)
    os.remove(zip_path)

    return response



def health_check(request):
    # You can include additional health checks here
    return JsonResponse({"status": "ok", "message": "Service is operational"}, status=200)